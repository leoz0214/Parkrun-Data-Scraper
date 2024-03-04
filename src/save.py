"""
Module for handling exporting data to various file types:
CSV, XLSX, DOCX and PDF.
"""
import csv
import io
import pathlib
import platform
import tempfile
from contextlib import suppress

try:
    import comtypes.client as comtypesclient
except ImportError:
    comtypesclient = None
import docx
import docx.document
import docx.table
with suppress(ImportError):
    import docx2pdf
import openpyxl
from docx.shared import Inches, Pt
from matplotlib import pyplot as plt
from openpyxl.styles.fonts import Font

import output
from utils import seconds_to_mmss


# Tabular output columns
TABLE_COLUMNS = (
    "Event Number", "Date", "Finishers", "Volunteers",
    "Male 1st Name", "Male 1st Athlete ID", "Male 1st Seconds",
    "Female 1st Name", "Female 1st Athlete ID", "Female 1st Seconds"
)
MAX_SHEET_NAME_LENGTH = 31
GRAPH_WIDTH = Inches(5)
GRAPH_HEIGHT = Inches(3)
TITLE_SIZE = Pt(28)
HEADING_1_SIZE = Pt(24)
TEXT_SIZE = Pt(16)
TOP_WINNERS_COLUMNS_AND_WIDTHS = {
    "#": Inches(0.25), "Name": Inches(2.25),
    "Athlete ID": Inches(0.9), "Wins": Inches(0.4)
}
TABLE_SIZE = Pt(11)
WORD_EXPORT_PDF_ID = 17


def get_event_records(events: list["output.EventData"]) -> list[list]:
    """Converts event data objects into records to be written to CSV/XLSX."""
    records = []
    for event in events:
        record = [event.number, event.date, event.finishers, event.volunteers]
        for first in (event.first_male, event.first_female):
            if first is not None:
                record.extend((first.name, first.athlete_id, first.seconds))
            else:
                record.extend((None, None, None))
        records.append(record)
    return records


def save_csv(events: list["output.EventData"], file_path: str) -> None:
    """Saves event data to CSV."""
    records = get_event_records(events)
    with open(file_path, "w", encoding="utf8") as f:
        writer = csv.writer(f, lineterminator="\n")
        writer.writerow(TABLE_COLUMNS)
        writer.writerows(records)


def save_xlsx(data: "output.Data", file_path: str) -> None:
    """Saves event data to XLSX."""
    records = get_event_records(data.events)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = f"{data.title} Parkrun"[:MAX_SHEET_NAME_LENGTH]
    sheet.append(TABLE_COLUMNS)
    # Make headings bold.
    for column in range(1, len(TABLE_COLUMNS) + 1):
        sheet.cell(row=1, column=column).font = Font(bold=True)
    for record in records:
        sheet.append(record)
    workbook.save(file_path)


def add_graph(
    document: docx.document.Document, data: "output.Data",
    label: str, attribute: str
) -> None:
    """Adds graph for given metric to given document."""
    plt.figure()
    output.plot_graph(data, label, attribute)
    with io.BytesIO() as graph_bytes:
        plt.savefig(graph_bytes)
        graph_bytes.seek(0)
        document.add_picture(graph_bytes, GRAPH_WIDTH, GRAPH_HEIGHT)
    plt.close()


def set_table_font_size(table: docx.table.Table, size: Pt) -> None:
    """Sets the font size of the entire table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = size


def set_table_column_widths(table: docx.table.Table, widths) -> None:
    """Sets the column widths of table."""
    for column, width in zip(table.columns, widths):
        for cell in column.cells:
            cell.width = width


def add_top_winners_table(
    document: docx.document.Document, top_winners: list["output.TopWinner"]
) -> None:
    """Creates top male/female winners table."""
    table = document.add_table(
        rows=len(top_winners) + 1,
        cols=len(TOP_WINNERS_COLUMNS_AND_WIDTHS))
    table.style = "Table Grid"
    table.autofit = True
    for i, column in enumerate(TOP_WINNERS_COLUMNS_AND_WIDTHS):
        cell = table.cell(0, i)
        cell.text = str(column)
        # Make cell (column) text bold.
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for i, row in enumerate(top_winners, 1):
        record = (i, row.name, row.athlete_id, row.wins)
        for j, value in enumerate(record):
            table.cell(i, j).text = str(value)
    set_table_font_size(table, TABLE_SIZE)
    set_table_column_widths(table, TOP_WINNERS_COLUMNS_AND_WIDTHS.values())


def generate_docx(data: "output.Data") -> docx.document.Document:
    """Generates and returns a DOCX report ready to be saved."""
    document: docx.document.Document = docx.Document()
    document.styles["Title"].font.size = TITLE_SIZE
    document.styles["Heading 1"].font.size = HEADING_1_SIZE
    document.styles["Normal"].font.size = TEXT_SIZE
    document.add_heading(f"{data.title} Parkrun Stats", 0)
    document.add_heading("Event Popularity", 1)
    popularity_points = (
        f"Mean finishers: {data.mean_finishers:.1f}",
        f"Median finishers: {data.median_finishers}",
        f"Mean volunteers: {data.mean_volunteers:.1f}",
        f"Median volunteers: {data.median_volunteers}")
    for point in popularity_points:
        document.add_paragraph(point, style="List Bullet")
    add_graph(document, data, "Finishers", "finishers")
    add_graph(document, data, "Volunteers", "volunteers")

    document.add_heading("Competitive", 1)
    document.add_paragraph(
        f"The male course record is held by {data.male_record.name} "
        f"(A{data.male_record.athlete_id}), with a time of "
        f"{seconds_to_mmss(data.male_record.seconds)}.")
    document.add_paragraph(
        f"The female course record is held by {data.female_record.name} "
        f"(A{data.female_record.athlete_id}), with a time of "
        f"{seconds_to_mmss(data.female_record.seconds)}.")
    document.add_paragraph(
        f"The mean male 1st place time is "
        f"{seconds_to_mmss(data.mean_first_male_seconds)}.")
    document.add_paragraph(
        f"The mean female 1st place time is: "
        f"{seconds_to_mmss(data.mean_first_female_seconds)}.")
    add_graph(document, data, "Male 1st Time", "first_male.seconds")
    add_graph(document, data, "Female 1st Time", "first_female.seconds")
    document.add_paragraph("Top male winners:")
    add_top_winners_table(document, data.top_male_winners)
    document.add_paragraph()
    document.add_paragraph("Top female winners:")
    add_top_winners_table(document, data.top_female_winners)
    document.add_paragraph()

    document.add_heading("Summary", 1)
    summary_points = (
        f"Event count: {data.event_count}",
        f"Finishes: {data.finishes}", f"Finishers: {data.finishers}",
        f"Volunteers: {data.volunteers}", f"Groups: {data.groups}",
        f"Personal bests: {data.personal_bests}",
        f"Mean finish time: {seconds_to_mmss(data.mean_seconds)}",
        f"Email: {data.email}")
    for point in summary_points:
        document.add_paragraph(point, style="List Bullet")

    return document


def save_docx(data: "output.Data", file_path: str) -> None:
    """Saves a DOCX report on event data."""
    document = generate_docx(data)
    document.save(file_path)


def save_pdf(data: "output.Data", file_path: str) -> None:
    """Saves a PDF report on event data."""
    document = generate_docx(data)
    # Create temp file to use as dummy DOCX.
    with tempfile.NamedTemporaryFile("wb", suffix=".docx", delete=False) as f:
        temp_file_path = pathlib.Path(f.name)
        document.save(f)
    try:
        if platform.system() == "Windows" and comtypesclient is not None:
            # Use MS Word on windows to attempt conversion.
            # Avoids weird behaviour of docx2pdf.
            word = None
            doc = None
            try:
                word = comtypesclient.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(temp_file_path))
                doc.SaveAs(file_path, FileFormat=WORD_EXPORT_PDF_ID)
            finally:
                if doc is not None:
                    doc.Close()
                if word is not None:
                    word.Quit()
        else:
            # For MacOS, attempt conversion using docx2pdf.
            docx2pdf.convert(temp_file_path, file_path, keep_active=True)
    finally:
        temp_file_path.unlink(missing_ok=True)
